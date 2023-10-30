#import basic libs
from pdf2docx import parse
import re
from docx import Document

#define function for main conversion
def convert_pdf2docx(input_file: str, output_file: str):
    """Converts pdf to docx"""
    
    result = parse(input_file, output_file)
    summary = {
        "File": input_file, "Output File": output_file
    }
    # Printing Summary
    print("## Summary ########################################################")
    print("\n".join("{}:{}".format(i, j) for i, j in summary.items()))
    print("###################################################################")
    return result

if __name__ == "__main__":
    input_file = input("Please enter the path to your PDF: ").strip('\"')
    interfile = input_file.split(".")
    output_file = interfile[0] + ".docx"
    convert_pdf2docx(input_file, output_file)

    # Fix artifacts from conversion
    doc = Document(output_file)
    for para in doc.paragraphs:
        #if you don't go all the way down to "run," you'll lose your formatting
        for run in para.runs:
            run.text = re.sub('!', ' ', run.text)
            run.text = re.sub(r'(?<=[a-zA-Z])\.(?=[a-zA-Z])', ' ', run.text)
            run.text = re.sub(r'\.\.(?=[A-Z])', '. ', run.text)
            run.text = re.sub(r'(?<=\w)\((?=\w)', ' ', run.text)
            run.text = re.sub(r'^\.$', '', run.text)
            run.text = re.sub(r'\.$', '', run.text)
            run.text = re.sub(r'\.\.\.+', '.', run.text)
            run.text = re.sub(r',\.', ', ', run.text)
            run.text = re.sub(r'  ', ' ', run.text)
    # have to drill down a bit to get to the cell content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = re.sub(r'(?<=[a-zA-Z])\.|!+(?=$)', ' ', cell.text)
                cell.text = re.sub(r'(?<=[a-zA-Z])!(?=[a-zA-Z])', ' ', cell.text)

    doc.save(output_file)
